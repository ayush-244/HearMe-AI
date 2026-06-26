import logging
from pathlib import Path
from .common import ExtractedDocument, DocumentNormalizer

logger = logging.getLogger(__name__)


class DOCXLoader:
    def extract(self, filepath: str) -> ExtractedDocument:
        path = Path(filepath)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {filepath}")

        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required. Install: pip install python-docx")

        try:
            doc = docx.Document(filepath)
        except Exception as e:
            logger.error("DOCX extraction failed (corrupted): %s — %s", filepath, e)
            raise ValueError(f"Corrupted DOCX: {e}")

        text_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        raw_text = "\n\n".join(text_parts)
        text = DocumentNormalizer.normalize(raw_text)

        core_properties = {}
        try:
            props = doc.core_properties
            for attr in ("title", "author", "subject", "keywords", "category", "comments"):
                val = getattr(props, attr, None)
                if val:
                    core_properties[attr] = str(val)
        except Exception:
            pass

        preview = DocumentNormalizer.generate_preview(text)
        words = DocumentNormalizer.count_words(text)
        characters = DocumentNormalizer.count_characters(text)

        logger.info(
            "DOCX extracted: words=%d, chars=%d",
            words, characters,
        )

        return ExtractedDocument(
            text=text,
            preview=preview,
            pages=0,
            words=words,
            characters=characters,
            metadata=core_properties,
        )
