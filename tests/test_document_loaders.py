"""Unit tests for document loaders."""
import pytest
from pathlib import Path
from ai.documents.pdf_loader import PDFLoader
from ai.documents.docx_loader import DOCXLoader
from ai.documents.txt_loader import TXTLoader
from ai.documents.markdown_loader import MarkdownLoader
from ai.documents.common import ExtractedDocument


class TestPDFLoader:
    @pytest.fixture
    def loader(self):
        return PDFLoader()

    @pytest.fixture
    def sample_pdf(self, tmp_path):
        import fitz
        path = tmp_path / "sample.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello World", fontsize=12)
        page.insert_text((72, 100), "Page 1 content", fontsize=12)
        doc.save(str(path))
        doc.close()
        return str(path)

    @pytest.fixture
    def multi_page_pdf(self, tmp_path):
        import fitz
        path = tmp_path / "multi.pdf"
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content", fontsize=12)
        doc.save(str(path))
        doc.close()
        return str(path)

    @pytest.fixture
    def password_pdf(self, tmp_path):
        import fitz
        path = tmp_path / "protected.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Secret", fontsize=12)
        doc.save(str(path), encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="password", owner_pw="password")
        doc.close()
        return str(path)

    def test_extract_text(self, loader, sample_pdf):
        result = loader.extract(sample_pdf)
        assert isinstance(result, ExtractedDocument)
        assert "Hello World" in result.text
        assert result.pages == 1
        assert result.words > 0
        assert result.characters > 0
        assert len(result.preview) > 0

    def test_multi_page(self, loader, multi_page_pdf):
        result = loader.extract(multi_page_pdf)
        assert result.pages == 3
        assert "Page 1" in result.text
        assert "Page 2" in result.text
        assert "Page 3" in result.text

    def test_password_protected(self, loader, password_pdf):
        with pytest.raises(ValueError, match="password"):
            loader.extract(password_pdf)

    def test_corrupted_pdf(self, loader, tmp_path):
        path = tmp_path / "corrupted.pdf"
        path.write_bytes(b"This is not a valid PDF file at all")
        with pytest.raises(ValueError, match="Corrupted PDF"):
            loader.extract(str(path))

    def test_file_not_found(self, loader):
        with pytest.raises(FileNotFoundError):
            loader.extract("/nonexistent/path.pdf")

    def test_empty_pdf(self, loader, tmp_path):
        import fitz
        path = tmp_path / "empty.pdf"
        doc = fitz.open()
        doc.new_page()
        doc.save(str(path))
        doc.close()
        result = loader.extract(str(path))
        assert result.pages == 1
        assert result.text == ""
        assert result.words == 0

    def test_metadata_preserved(self, loader, tmp_path):
        import fitz
        path = tmp_path / "meta.pdf"
        doc = fitz.open()
        doc.set_metadata({"title": "Test Doc", "author": "Tester", "subject": "Testing"})
        page = doc.new_page()
        page.insert_text((72, 72), "Metadata test", fontsize=12)
        doc.save(str(path))
        doc.close()
        result = loader.extract(str(path))
        assert result.metadata.get("title") == "Test Doc"
        assert result.metadata.get("author") == "Tester"

    def test_preview_is_truncated(self, loader, tmp_path):
        import fitz
        path = tmp_path / "long.pdf"
        doc = fitz.open()
        for _ in range(10):
            page = doc.new_page()
            page.insert_text((72, 72), "Lorem ipsum dolor sit amet consectetur adipiscing elit sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. ", fontsize=12)
        doc.save(str(path))
        doc.close()
        result = loader.extract(str(path))
        assert result.preview.endswith("...")


class TestDOCXLoader:
    @pytest.fixture
    def loader(self):
        return DOCXLoader()

    @pytest.fixture
    def sample_docx(self, tmp_path):
        from docx import Document
        path = tmp_path / "sample.docx"
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Second paragraph with more text.")
        doc.save(str(path))
        return str(path)

    @pytest.fixture
    def docx_with_table(self, tmp_path):
        from docx import Document
        path = tmp_path / "table.docx"
        doc = Document()
        doc.add_paragraph("Text before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Item1"
        table.cell(1, 1).text = "42"
        doc.add_paragraph("Text after table")
        doc.save(str(path))
        return str(path)

    def test_extract_text(self, loader, sample_docx):
        result = loader.extract(sample_docx)
        assert isinstance(result, ExtractedDocument)
        assert "Hello World" in result.text
        assert "Second paragraph" in result.text

    def test_table_flattened(self, loader, docx_with_table):
        result = loader.extract(docx_with_table)
        assert "Name" in result.text
        assert "Value" in result.text
        assert "Item1" in result.text
        assert "42" in result.text
        assert "Text before table" in result.text

    def test_file_not_found(self, loader):
        with pytest.raises(FileNotFoundError):
            loader.extract("/nonexistent.docx")

    def test_corrupted_docx(self, loader, tmp_path):
        path = tmp_path / "corrupted.docx"
        path.write_bytes(b"not a zip file")
        with pytest.raises(ValueError, match="Corrupted DOCX"):
            loader.extract(str(path))


class TestTXTLoader:
    @pytest.fixture
    def loader(self):
        return TXTLoader()

    def test_utf8(self, loader, tmp_path):
        path = tmp_path / "test.txt"
        path.write_text("Hello World\nSecond line", encoding="utf-8")
        result = loader.extract(str(path))
        assert "Hello World" in result.text
        assert "Second line" in result.text

    def test_utf16_le(self, loader, tmp_path):
        path = tmp_path / "utf16.txt"
        path.write_bytes("Hello UTF-16".encode("utf-16-le"))
        result = loader.extract(str(path))
        assert "Hello UTF-16" in result.text

    def test_utf16_be(self, loader, tmp_path):
        path = tmp_path / "utf16be.txt"
        path.write_bytes("Hello BE".encode("utf-16-be"))
        result = loader.extract(str(path))
        assert "Hello BE" in result.text

    def test_latin1_fallback(self, loader, tmp_path):
        path = tmp_path / "latin1.txt"
        path.write_bytes("Latin-1 text: \xe9\xe0\xfc".encode("latin-1"))
        result = loader.extract(str(path))
        assert "Latin-1" in result.text

    def test_file_not_found(self, loader):
        with pytest.raises(FileNotFoundError):
            loader.extract("/nonexistent.txt")

    def test_encoding_fallback_logs_warning(self, loader, tmp_path):
        path = tmp_path / "binary.txt"
        path.write_bytes(b"\xff\xfe\x00\xff\x00\xfe\x00\xff")
        result = loader.extract(str(path))
        assert isinstance(result, ExtractedDocument)


class TestMarkdownLoader:
    @pytest.fixture
    def loader(self):
        return MarkdownLoader()

    def test_bold_italic_stripped(self, loader, tmp_path):
        path = tmp_path / "test.md"
        path.write_text("**bold** and *italic* and `code`", encoding="utf-8")
        result = loader.extract(str(path))
        assert "bold" in result.text
        assert "italic" in result.text
        assert "code" in result.text
        assert "**" not in result.text

    def test_links_stripped(self, loader, tmp_path):
        path = tmp_path / "links.md"
        path.write_text("[click here](https://example.com)", encoding="utf-8")
        result = loader.extract(str(path))
        assert "click here" in result.text
        assert "https" not in result.text

    def test_headings_stripped(self, loader, tmp_path):
        path = tmp_path / "headings.md"
        path.write_text("# Title\n## Subtitle\n### Section", encoding="utf-8")
        result = loader.extract(str(path))
        assert "Title" in result.text
        assert "Subtitle" in result.text
        assert "Section" in result.text
        assert "#" not in result.text

    def test_code_blocks_stripped(self, loader, tmp_path):
        path = tmp_path / "code.md"
        path.write_text("Text\n```python\nprint('hello')\n```\nMore text", encoding="utf-8")
        result = loader.extract(str(path))
        assert "Text" in result.text
        assert "More text" in result.text
        assert "print" not in result.text

    def test_lists_stripped(self, loader, tmp_path):
        path = tmp_path / "list.md"
        path.write_text("- Item 1\n- Item 2\n- Item 3", encoding="utf-8")
        result = loader.extract(str(path))
        assert "Item 1" in result.text
        assert "Item 2" in result.text
        assert "Item 3" in result.text
        assert "-" not in result.text

    def test_blockquotes_stripped(self, loader, tmp_path):
        path = tmp_path / "quote.md"
        path.write_text("> This is a quote\n> More quote", encoding="utf-8")
        result = loader.extract(str(path))
        assert "This is a quote" in result.text
        assert ">" not in result.text

    def test_file_not_found(self, loader):
        with pytest.raises(FileNotFoundError):
            loader.extract("/nonexistent.md")


class TestDocumentNormalizer:
    def test_normalize_unicode(self):
        from ai.documents.common import DocumentNormalizer
        text = "\u201cHello\u201d"
        result = DocumentNormalizer.normalize(text)
        assert "Hello" in result

    def test_normalize_line_endings(self):
        from ai.documents.common import DocumentNormalizer
        text = "line1\r\nline2\rline3\nline4"
        result = DocumentNormalizer.normalize(text)
        assert "\r\n" not in result
        assert "\r" not in result

    def test_normalize_excessive_newlines(self):
        from ai.documents.common import DocumentNormalizer
        text = "para1\n\n\n\n\npara2"
        result = DocumentNormalizer.normalize(text)
        assert result == "para1\n\npara2"

    def test_generate_preview_short(self):
        from ai.documents.common import DocumentNormalizer
        text = "short text"
        preview = DocumentNormalizer.generate_preview(text, max_chars=500)
        assert preview == "short text"

    def test_generate_preview_long(self):
        from ai.documents.common import DocumentNormalizer
        text = "word " * 200
        preview = DocumentNormalizer.generate_preview(text, max_chars=100)
        assert len(preview) <= 110
        assert preview.endswith("...")

    def test_count_words(self):
        from ai.documents.common import DocumentNormalizer
        assert DocumentNormalizer.count_words("one two three") == 3
        assert DocumentNormalizer.count_words("") == 0

    def test_count_characters(self):
        from ai.documents.common import DocumentNormalizer
        assert DocumentNormalizer.count_characters("hello") == 5
        assert DocumentNormalizer.count_characters("") == 0
